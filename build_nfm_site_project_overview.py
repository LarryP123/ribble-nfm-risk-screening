from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "Natural_Flood_Management_Candidate_Sites_Project_Overview.docx"
NAVY, BLUE, TEAL = "17324D", "2E74B5", "247D78"
INK, MUTED, WHITE = "243142", "667085", "FFFFFF"
LIGHT, PALE_BLUE, PALE_TEAL, GOLD = "F2F4F7", "E8EEF5", "E6F2F0", "8A6500"


def font(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold, run.italic = bold, italic


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    node = pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        pr.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=90, bottom=90, start=120, end=120):
    pr = cell._tc.get_or_add_tcPr()
    tc_mar = pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        el = tc_mar.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def geometry(table, widths):
    table.autofit = False
    pr = table._tbl.tblPr
    tw = pr.find(qn("w:tblW"))
    if tw is None:
        tw = OxmlElement("w:tblW")
        pr.append(tw)
    tw.set(qn("w:w"), str(sum(widths)))
    tw.set(qn("w:type"), "dxa")
    ti = pr.find(qn("w:tblInd"))
    if ti is None:
        ti = OxmlElement("w:tblInd")
        pr.append(ti)
    ti.set(qn("w:w"), "120")
    ti.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            pr = cell._tc.get_or_add_tcPr()
            tcw = pr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                pr.append(tcw)
            tcw.set(qn("w:w"), str(widths[i]))
            tcw.set(qn("w:type"), "dxa")
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    pr.append(el)


def table(doc, headers, rows, widths, size=8.9, header_fill=NAVY):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    repeat_header(t.rows[0])
    for i, text in enumerate(headers):
        shade(t.rows[0].cells[i], header_fill)
        p = t.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(text), size, WHITE, True)
    for ri, values in enumerate(rows):
        cells = t.add_row().cells
        if ri % 2:
            for c in cells:
                shade(c, LIGHT)
        for i, text in enumerate(values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(str(text)), size)
    geometry(t, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    return t


def body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(" " + text)
    return p


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def callout(doc, label, text, fill=PALE_TEAL, accent=TEAL):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    geometry(t, [9360])
    c = t.cell(0, 0)
    shade(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(label.upper()), 9, accent, True)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    font(p2.add_run(text), 11.3, NAVY, True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

hp = sec.header.paragraphs[0]
hp.text = "NATURAL FLOOD MANAGEMENT  |  CANDIDATE SITE PROJECT"
hp.paragraph_format.space_after = Pt(0)
for r in hp.runs:
    font(r, 8.5, MUTED, True)

fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(fp.add_run("Project overview  •  "), 8.5, MUTED)
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
fp._p.append(field)

styles = doc.styles
n = styles["Normal"]
n.font.name, n.font.size = "Calibri", Pt(11)
n.font.color.rgb = RGBColor.from_string(INK)
n.paragraph_format.space_before = Pt(0)
n.paragraph_format.space_after = Pt(6)
n.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, NAVY, 8, 4),
]:
    s = styles[name]
    s.font.name, s.font.size, s.font.bold = "Calibri", Pt(size), True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(before), Pt(after)
    s.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    s = styles[name]
    s.font.name, s.font.size = "Calibri", Pt(11)
    s.paragraph_format.space_after = Pt(8)
    s.paragraph_format.line_spacing = 1.167

# First-page memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(7)
font(p.add_run("GEOSPATIAL RISK PORTFOLIO PROJECT"), 10, TEAL, True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
font(p.add_run("Finding Overlooked Natural\nFlood-Management Sites"), 27, NAVY, True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(15)
font(p.add_run("A data-led screening and local-validation project for a Rivers Trust or council flood team"), 13.5, MUTED)

table(doc, ["PRIMARY USER", "DECISION", "INITIAL SCOPE"], [[
    "One local Rivers Trust",
    "Which areas merit investigation?",
    "One catchment; five sites",
]], [3000, 3600, 2760], 9.2, TEAL)

callout(doc, "Research question", "Which locations show strong potential for natural flood management and habitat recovery, but little recorded delivery—and are they genuine opportunities or gaps in the public data?")

doc.add_heading("1. Purpose", level=1)
body(doc, "The project will combine national flood, habitat and restoration datasets to identify candidate areas for natural flood management within one selected catchment. It will consider several intervention types—including floodplain reconnection, wet woodland, ponds, runoff storage and land-management measures—rather than assuming wetlands are always the answer.")
body(doc, "The initial shortlist will then be tested with a local Rivers Trust or council officer. Their feedback will help distinguish genuinely overlooked opportunities from locations affected by unrecorded work, land constraints or limitations in the national model.")

doc.add_heading("Intended outcome", level=2)
body(doc, "A credible, visually clear briefing that identifies five candidate investigation areas, explains the evidence and uncertainty behind each one, and shows how professional challenge changed the final recommendation.")

doc.add_page_break()
doc.add_heading("2. Scope and analytical question", level=1)

doc.add_heading("Study geography", level=2)
body(doc, "Select one catchment with an active Rivers Trust, identifiable downstream communities and accessible spatial data. A focused geography makes map production, manual review and stakeholder validation manageable.")

doc.add_heading("What the project is looking for", level=2)
for text in [
    "High modelled potential for one or more natural flood-management measures.",
    "A plausible connection to downstream people, property or infrastructure exposed to flooding.",
    "An opportunity to restore, expand or connect priority habitat.",
    "Little overlap with recorded restoration or natural flood-management activity.",
    "No obvious major constraint visible in the available public data.",
    "Enough evidence to justify local review or a more detailed feasibility assessment.",
]:
    bullet(doc, text)

doc.add_heading("What the project will not claim", level=2)
for text in [
    "That a candidate area is ready for construction or investment.",
    "That public data provides complete information about local projects or land ownership.",
    "That a screening score replaces hydrological, ecological or engineering assessment.",
    "That a particular intervention will produce a guaranteed reduction in flood losses.",
]:
    bullet(doc, text)

callout(doc, "Important distinction", "The output identifies candidate investigation areas—not confirmed project sites. Its value is in narrowing the search and making uncertainty visible.", PALE_BLUE, BLUE)

doc.add_page_break()
doc.add_heading("3. Data and workflow", level=1)

table(doc, ["DATA", "PURPOSE", "CORE CONTROL"], [
    ("EA NFM Heat Maps", "Locate nationally modelled opportunities by intervention type", "Record model date, resolution and interpretation limits"),
    ("Flood risk and historic flooding", "Identify downstream receptors and recurring problems", "Avoid treating risk bands as property-level predictions"),
    ("Catchments, rivers and terrain", "Establish hydrological and upstream/downstream context", "Use a consistent CRS and validate topology"),
    ("Priority Habitats Inventory", "Identify existing habitat and restoration context", "Check habitat classes and currency"),
    ("Habitat Networks", "Find enhancement, expansion and connectivity opportunities", "Separate mapped potential from verified condition"),
    ("Recorded restoration projects", "Test whether an opportunity already has recorded activity", "Treat absence as uncertain, not proof of inactivity"),
    ("Land cover, roads and settlements", "Support site interpretation and identify visible constraints", "Do not infer ownership or consent"),
], [2400, 4200, 2760], 8.6)

doc.add_heading("Workflow", level=2)
steps = [
    "Choose a catchment and define the spatial unit used for screening.",
    "Download each source, retain the raw files and create a provenance manifest.",
    "Standardise coordinate systems, schemas and geographic identifiers.",
    "Calculate flood potential, downstream exposure, habitat opportunity and recorded activity.",
    "Remove or flag locations with obvious constraints and low data confidence.",
    "Rank the remaining areas and manually review the strongest candidates.",
    "Create five site cards and send them to a local practitioner for validation.",
    "Revise the shortlist and document why candidates were retained, downgraded or rejected.",
]
for step in steps:
    number(doc, step)

doc.add_page_break()
doc.add_heading("4. Candidate screening and maps", level=1)

doc.add_heading("Transparent screening framework", level=2)
table(doc, ["FACTOR", "QUESTION", "EFFECT"], [
    ("NFM potential", "Does the national model indicate a strong opportunity?", "Increase score"),
    ("Downstream exposure", "Are communities, roads or assets plausibly downstream?", "Increase score"),
    ("Habitat opportunity", "Could action restore or connect habitat?", "Increase score"),
    ("Recorded activity", "Is restoration or NFM already recorded nearby?", "Reduce gap score"),
    ("Visible constraints", "Do settlement, infrastructure or protected features limit delivery?", "Reduce or flag"),
    ("Data confidence", "Are sources current, complete and spatially consistent?", "Qualify recommendation"),
], [2000, 5000, 2360], 8.8)

body(doc, "Weights will be explicit and tested through sensitivity analysis. Candidate ranks should be reported in bands, with rank stability and confidence, rather than presented as precise truth.")

doc.add_heading("Candidate site maps", level=2)
for text in [
    "Catchment overview: all candidate areas, main rivers and downstream settlements.",
    "Opportunity-gap map: high NFM potential compared with recorded activity.",
    "Five compact site maps: satellite basemap, shaded opportunity areas, suggested rewetting zones, rivers, floodplain, habitat layers, labels, north arrow, scale and source note.",
    "Locator inset: each site’s position within the wider catchment.",
    "Validation map: retained, rejected and already-known sites after practitioner review.",
]:
    bullet(doc, text)

doc.add_heading("Site-card contents", level=2)
table(doc, ["MAP PANEL", "EVIDENCE PANEL"], [[
    "Compact site map using clear shading to distinguish stronger opportunity areas and locations where rewetting could be investigated.",
    "Location; possible interventions; downstream exposure; habitat value; recorded activity; constraints; confidence; recommended next check.",
]], [4300, 5060], 9.2, TEAL)

doc.add_page_break()
doc.add_heading("5. Deliverables and definition of done", level=1)

table(doc, ["ID", "DELIVERABLE", "DEFINITION OF DONE"], [
    ("D1", "Reproducible spatial pipeline", "A documented command rebuilds the analytical outputs from retained source files."),
    ("D2", "Data dictionary and provenance manifest", "Every field, source, date, licence and major transformation is recorded."),
    ("D3", "Catchment opportunity map", "All screened locations are visible with an intelligible legend and confidence layer."),
    ("D4", "Five candidate site cards", "Each card contains a small site map with shaded opportunity and suggested rewetting areas, plus evidence, constraints and next action."),
    ("D5", "Interactive map", "A user can inspect candidates, switch core layers and trace each score to its inputs."),
    ("D6", "Validation record", "Practitioner feedback and its effect on each candidate are documented."),
    ("D7", "Insight briefing", "A concise report states the findings, limitations and recommended feasibility checks."),
    ("D8", "Portfolio case study", "Architecture, controls, maps, findings and reflection are ready for an application or interview."),
], [620, 2600, 6140], 8.7)

doc.add_heading("Suggested final finding structure", level=2)
callout(doc, "Example—not a predetermined conclusion", "The national data identified eight apparent opportunity gaps. Local review found that three already had unrecorded activity, two faced significant delivery constraints, and three remained credible candidates for feasibility work.", PALE_BLUE, BLUE)

doc.add_heading("Success measures", level=2)
for text in [
    "At least five plausible candidates are assessed consistently.",
    "Every recommendation can be traced to source data and explicit judgement.",
    "The shortlist remains reasonably stable under alternative weights.",
    "At least one local practitioner reviews the candidate maps.",
    "The final case study explains false positives and data gaps, not only successful candidates.",
]:
    bullet(doc, text)

doc.add_page_break()
doc.add_heading("6. Delivery plan and outreach", level=1)
table(doc, ["WEEK", "FOCUS", "OUTPUT"], [
    ("1", "Select catchment; acquire and catalogue data", "Scope note and source manifest"),
    ("2", "Build validated spatial staging layers", "Tested spatial database"),
    ("3", "Create screening features and initial ranking", "Candidate longlist"),
    ("4", "Manual review and map production", "Five draft site cards"),
    ("5", "Practitioner outreach and validation", "Feedback and decision log"),
    ("6", "Revise analysis and publish case study", "Final maps, briefing and repository"),
], [800, 5000, 3560], 8.9)

doc.add_heading("Outreach approach", level=2)
body(doc, "Approach one named catchment officer, natural flood-management lead or local authority flood officer with a single-page summary and one strong map. Ask for technical feedback rather than funding or endorsement.")

callout(doc, "Suggested message", "I used the Environment Agency’s new NFM opportunity data alongside habitat and recorded restoration layers to identify five possible gaps in your catchment. Could I show you the maps and learn whether they reflect local conditions or reveal limitations in the public data?")

doc.add_heading("Starting sources", level=2)
sources = [
    "Environment Agency — Natural Flood Management Heat Maps (published 2026): https://environment.data.gov.uk/dataset/ecc7d246-ae51-46ac-862f-11da22b72366",
    "Natural England — Priority Habitats Inventory: https://www.data.gov.uk/dataset/4b6ddab7-6c0f-4407-946e-d6499f19fcde/priority-habitats-inventory-england",
    "Natural England — Habitat Networks (Combined Habitats): https://www.data.gov.uk/dataset/0ef2ed26-2f04-4e0f-9493-ffbdbfaeb159/habitat-networks-combined-habitats-england",
    "Environment Agency — Priority Habitat Creation and Restoration: https://www.data.gov.uk/dataset/e0165747-8368-4ff7-a644-df9aeb27bb0b/priority-habitat-creation-and-restoration",
]
for source in sources:
    p = bullet(doc, source)
    for r in p.runs:
        font(r, 9, MUTED)

doc.core_properties.title = "Finding Overlooked Natural Flood-Management Sites"
doc.core_properties.subject = "Geospatial risk portfolio project overview"
doc.core_properties.author = "Portfolio project"
doc.core_properties.keywords = "natural flood management, geospatial, risk analysis, habitat recovery, candidate sites"
doc.save(OUT)
print(OUT)
