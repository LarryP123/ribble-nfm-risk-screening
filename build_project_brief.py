from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_ROW_HEIGHT_RULE


OUT = "Wetlands_Flood_Risk_Insurance_Project_Brief.docx"
NAVY = "16324F"
BLUE = "2E74B5"
TEAL = "207F7A"
PALE_TEAL = "E7F3F1"
PALE_BLUE = "E8EEF5"
LIGHT = "F2F4F7"
MID = "667085"
INK = "1F2937"
WHITE = "FFFFFF"
GOLD = "9A6700"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_keep_with_next(p, keep=True):
    p.paragraph_format.keep_with_next = keep


def set_repeat_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = "WETLANDS AS FLOOD-RISK INFRASTRUCTURE  |  PROJECT BRIEF"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(MID)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Portfolio project specification  •  ")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MID)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def add_title(doc, text, size=27, color=NAVY, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text.upper())
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    r.font.letter_spacing = Pt(0.5) if hasattr(r.font, "letter_spacing") else None
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        r1.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, label, text, fill=PALE_TEAL, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    rr = p2.add_run(text)
    rr.font.size = Pt(11.5)
    rr.bold = True
    rr.font.color.rgb = RGBColor.from_string(NAVY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def add_table(doc, headers, rows, widths_dxa, header_fill=NAVY, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor.from_string(WHITE if header_fill == NAVY else NAVY)
    for row_i, values in enumerate(rows):
        cells = table.add_row().cells
        if row_i % 2 == 1:
            for c in cells:
                set_cell_shading(c, LIGHT)
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
            r.font.color.rgb = RGBColor.from_string(INK)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.86)
section.bottom_margin = Inches(0.78)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.42)
section.footer_distance = Inches(0.42)
set_repeat_header_footer(section)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(5.5)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 15, 7),
    ("Heading 2", 13, BLUE, 11, 5),
    ("Heading 3", 11.5, NAVY, 8, 3),
]:
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

for name in ["List Bullet", "List Bullet 2", "List Number"]:
    s = styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)
    s.paragraph_format.space_after = Pt(3.5)
    s.paragraph_format.line_spacing = 1.10

# Opening block
doc.add_paragraph().paragraph_format.space_after = Pt(7)
add_kicker(doc, "Risk analytics portfolio project")
add_title(doc, "Quantifying the Insurance Value\nof Wetland Recovery")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
r = p.add_run("A geospatial flood-loss and investment pipeline for a fictional UK home insurer")
r.font.name = "Calibri"
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string(MID)

meta = add_table(doc, ["AUDIENCE", "DECISION", "DELIVERY WINDOW"], [[
    "Climate Risk Committee",
    "Where should £10m be invested?",
    "Six weeks",
]], [3120, 3120, 3120], header_fill=TEAL, font_size=9.5)

add_callout(doc, "Project question", "Where could wetland recovery reduce household flood exposure and deliver the greatest risk-adjusted reduction in insured losses per pound invested?")

doc.add_heading("1. Project purpose", level=1)
add_body(doc, "This project will build a reproducible decision-support pipeline for a fictional UK home insurer considering investment in natural flood management. It will combine flood hazard, residential exposure, property vulnerability and wetland-restoration opportunity data to rank catchments for further investigation.")
add_body(doc, "The outcome is not an actuarially validated catastrophe model or a claim that wetlands prevent all flooding. It is a transparent screening model that helps a decision committee compare opportunities, understand uncertainty and decide where specialist feasibility work is justified.")

doc.add_heading("Primary audience", level=2)
add_body(doc, "The primary user is the insurer’s Climate Risk Committee, comprising representatives from portfolio management, catastrophe modelling, underwriting, sustainability and finance. The committee controls a hypothetical £10 million resilience fund and must approve a shortlist of catchments for detailed hydrological and commercial assessment.")

doc.add_heading("Decision to be supported", level=2)
add_body(doc, "The committee must select three to five catchments for feasibility studies and provide a documented rationale for funding, monitoring or rejecting each candidate.")

doc.add_page_break()
doc.add_heading("2. Objectives and success measures", level=1)
add_body(doc, "The project succeeds when it turns heterogeneous public data into a traceable financial-risk comparison that a non-technical committee can interrogate.")

objectives = [
    ("O1", "Create a repeatable ingestion pipeline", "At least four source datasets ingested with version, date, licence and provenance metadata."),
    ("O2", "Standardise spatial risk data", "All modelling tables use a documented catchment or grid geography and common coordinate reference system."),
    ("O3", "Estimate baseline flood loss", "Expected annual loss is calculated from explicit hazard, exposure and vulnerability assumptions."),
    ("O4", "Model wetland intervention scenarios", "Pessimistic, central and optimistic scenarios are available for every shortlisted catchment."),
    ("O5", "Rank investments transparently", "Each recommendation is reproducible from version-controlled weights, costs and assumptions."),
    ("O6", "Communicate uncertainty", "Every output carries a confidence rating, sensitivity result and limitations statement."),
]
add_table(doc, ["ID", "OBJECTIVE", "ACCEPTANCE MEASURE"], objectives, [720, 3000, 5640], font_size=9)

doc.add_heading("Out of scope", level=2)
for item in [
    "Predicting individual household claims or premium changes.",
    "Replacing detailed hydrological, ecological or engineering assessment.",
    "Using confidential insurer policy or claims records.",
    "Making causal claims that cannot be supported by the available data.",
    "Producing a production-grade capital model or regulatory submission.",
]:
    add_bullet(doc, item)

doc.add_heading("Portfolio evidence", level=2)
add_body(doc, "The finished project should demonstrate data engineering, spatial analytics, risk quantification, model governance, scenario analysis and communication to a senior decision-making audience.")

doc.add_page_break()
doc.add_heading("3. Analytical framework", level=1)

doc.add_heading("Core loss relationship", level=2)
add_callout(doc, "Financial measure", "Avoided Annual Loss = Baseline Expected Annual Loss − Post-restoration Expected Annual Loss", fill=PALE_BLUE, accent=BLUE)
add_body(doc, "At screening level, expected annual loss can be represented as the product of exposure, annual flood probability, vulnerability and insurance penetration. Each term must remain visible in the model rather than being hidden inside an unexplained composite score.")

add_table(doc, ["COMPONENT", "SCREENING REPRESENTATION", "KEY LIMITATION"], [
    ("Hazard", "Annual flood probability or risk band", "Public flood mapping is not a property-level claims model."),
    ("Exposure", "Residential properties and value proxy", "A synthetic book will not match a real insurer portfolio."),
    ("Vulnerability", "Damage ratio by depth/risk band", "Generic assumptions omit building-level characteristics."),
    ("Intervention", "Assumed loss-reduction range", "Effectiveness varies by catchment, storm and intervention design."),
    ("Cost", "Restored hectares × unit-cost range", "Land, maintenance and transaction costs may be incomplete."),
], [1550, 3430, 4380], font_size=8.8)

doc.add_heading("Required scenarios", level=2)
add_table(doc, ["SCENARIO", "PURPOSE", "ILLUSTRATIVE ASSUMPTION"], [
    ("Pessimistic", "Tests downside and weak intervention performance", "3% loss reduction; high cost; delayed benefits"),
    ("Central", "Provides the principal committee comparison", "8% loss reduction; central cost and delivery case"),
    ("Optimistic", "Tests upside without becoming the recommendation basis", "15% loss reduction; lower cost; timely delivery"),
], [1800, 3960, 3600], font_size=9)

doc.add_heading("Committee metrics", level=2)
for item in [
    "Baseline expected annual loss and modelled avoided annual loss.",
    "Risk-adjusted avoided loss per £1 million invested.",
    "Ten-year benefit-cost ratio and indicative payback period.",
    "Properties and population within the downstream exposure area.",
    "Restorable wetland hectares and habitat-connectivity potential.",
    "Confidence score and sensitivity of rank to key assumptions.",
]:
    add_bullet(doc, item)

doc.add_page_break()
doc.add_heading("4. Data and engineering specification", level=1)
add_body(doc, "Use public data and a clearly labelled synthetic insurance portfolio. The minimum viable product should prioritise dependable lineage and tests over the number of sources.")

add_table(doc, ["DATA DOMAIN", "EXAMPLE SOURCE", "USE IN MODEL", "CONTROL"], [
    ("Flood hazard", "Environment Agency flood-risk data", "Probability/risk bands and exposed areas", "Publication date, geometry and coverage checks"),
    ("Habitats", "Natural England Priority Habitats Inventory", "Existing wetland and habitat connectivity", "Schema, habitat class and duplicate checks"),
    ("Hydrology", "Catchment and river boundaries", "Common aggregation and upstream/downstream logic", "CRS and topology validation"),
    ("Property exposure", "Open residential counts and value proxy", "Synthetic sums insured and household exposure", "Range and completeness checks"),
    ("Recovery priorities", "Local Nature Recovery Strategy", "Strategic alignment and candidate intervention areas", "Version and geographic coverage"),
    ("Climate", "Rainfall or flood-uplift scenario", "Future-risk sensitivity", "Scenario label and no false precision"),
], [1620, 2520, 2940, 2280], font_size=8.2)

doc.add_heading("Target architecture", level=2)
for step in [
    "Extract: download or query source files while recording URL, retrieval date, licence and checksum.",
    "Stage: preserve immutable raw inputs and create a machine-readable ingestion manifest.",
    "Transform: standardise names, types, coordinate systems and spatial units in Python.",
    "Model: build tested analytical tables in DuckDB or PostgreSQL/PostGIS using dbt where practical.",
    "Score: generate scenario results, confidence ratings and recommendation classes.",
    "Serve: expose final tables through a Streamlit or Power BI decision dashboard.",
]:
    add_number(doc, step)

doc.add_heading("Minimum automated controls", level=2)
for item in [
    "Schema and required-column validation.",
    "Duplicate record and key uniqueness tests.",
    "Null thresholds for critical modelling fields.",
    "Invalid geometry, coordinate reference system and boundary checks.",
    "Freshness and source-version warnings.",
    "Reconciliation from source totals to analytical aggregates.",
    "Unit tests for expected-loss and scenario calculations.",
]:
    add_bullet(doc, item)

doc.add_page_break()
doc.add_heading("5. Deliverables", level=1)

deliverables = [
    ("D1", "Project repository", "Structured source code, configuration, licence notes and one-command setup.", "A new user can reproduce the sample output from the README."),
    ("D2", "Data pipeline", "Automated ingestion, standardisation and spatial transformation for the minimum datasets.", "Pipeline completes from clean environment; failures are actionable."),
    ("D3", "Analytical warehouse", "Documented raw, staging, feature and output tables.", "Keys, grain, lineage and field definitions are recorded."),
    ("D4", "Risk model", "Baseline loss, intervention scenarios, costs, avoided loss and confidence scores.", "All formulas and assumptions are version-controlled and tested."),
    ("D5", "Interactive dashboard", "Map, shortlist, comparison view and scenario controls for committee users.", "Users can explain why a location ranks where it does."),
    ("D6", "Committee paper", "Two-page recommendation with shortlist, rejected options, risks and next actions.", "Recommendations distinguish evidence, assumption and judgement."),
    ("D7", "Model card", "Intended use, exclusions, limitations, validation, monitoring and governance.", "No output can be mistaken for an actuarial or hydrological guarantee."),
    ("D8", "Portfolio case study", "Concise write-up with architecture, screenshots, findings and reflection.", "Suitable for GitHub and discussion in a risk-analyst interview."),
]
add_table(doc, ["ID", "DELIVERABLE", "CONTENTS", "DEFINITION OF DONE"], deliverables, [580, 1800, 3480, 3500], font_size=8.3)

doc.add_heading("Required dashboard views", level=2)
for item in [
    "National or regional overview with ranked catchments.",
    "Catchment detail showing hazard, exposure, wetland opportunity and assumptions.",
    "Scenario comparison for pessimistic, central and optimistic cases.",
    "Investment shortlist with avoided loss, cost, benefit-cost ratio and confidence.",
    "Data-quality and model-risk panel showing warnings and coverage gaps.",
]:
    add_bullet(doc, item)

doc.add_page_break()
doc.add_heading("6. Six-week delivery plan", level=1)

plan = [
    ("1", "Frame and acquire", "Select study geography; define model grain; download minimum data; create data dictionary and assumptions log.", "Source manifest and raw-data landing zone"),
    ("2", "Build the foundation", "Standardise spatial data; implement validation; create synthetic exposure portfolio.", "Tested staging tables"),
    ("3", "Estimate baseline risk", "Construct hazard × exposure × vulnerability logic and reconcile outputs.", "Baseline expected-loss table"),
    ("4", "Model interventions", "Create wetland-opportunity features; implement three scenarios, costs and confidence.", "Ranked scenario results"),
    ("5", "Build decision product", "Develop dashboard; test usability; draft committee recommendations.", "Dashboard beta and committee paper"),
    ("6", "Validate and present", "Run sensitivity tests; complete model card; polish README and case study; rehearse presentation.", "Portfolio-ready release"),
]
add_table(doc, ["WEEK", "FOCUS", "ACTIVITIES", "OUTPUT"], plan, [650, 1750, 4580, 2380], font_size=8.4)

doc.add_heading("Stage gates", level=2)
for item in [
    "Gate 1 — Scope: geography, audience, decision and minimum datasets agreed.",
    "Gate 2 — Data: ingestion is repeatable and critical quality tests pass.",
    "Gate 3 — Model: baseline loss reconciles and scenario assumptions are defensible.",
    "Gate 4 — Product: committee can reach the same recommendation from the evidence shown.",
    "Gate 5 — Release: limitations are prominent and the repository reproduces the published results.",
]:
    add_bullet(doc, item)

doc.add_heading("Minimum viable version", level=2)
add_body(doc, "If time is constrained, use one region, three to four source datasets, one common catchment geography and a synthetic portfolio. Deliver the central scenario plus sensitivity bounds, five candidate catchments, a simple interactive dashboard and a two-page committee paper.")

doc.add_page_break()
doc.add_heading("7. Risk and model governance", level=1)

risks = [
    ("R1", "Wetland benefit is overstated", "High", "Use ranges, cite assumptions, require hydrological validation before funding."),
    ("R2", "Synthetic exposure distorts priorities", "High", "Label synthetic results; test alternative portfolio distributions."),
    ("R3", "Public datasets have inconsistent dates", "Medium", "Store versions and retrieval dates; issue freshness warnings."),
    ("R4", "Spatial aggregation hides local variation", "Medium", "Document grain; compare results at an alternative aggregation where feasible."),
    ("R5", "Restoration cost proxy is incomplete", "High", "Include low/central/high costs and flag excluded lifecycle costs."),
    ("R6", "Benefits accrue to competing insurers", "Medium", "Treat co-funding and public-private partnership as an implementation condition."),
    ("R7", "Climate change reduces intervention performance", "High", "Run future-risk uplifts and do not assume static hazard."),
    ("R8", "Ranking creates false precision", "High", "Report bands, confidence and rank stability—not only a single ordered list."),
]
add_table(doc, ["ID", "RISK", "SIGNIFICANCE", "CONTROL / RESPONSE"], risks, [620, 2720, 1280, 4740], font_size=8.4)

doc.add_heading("Decision rules", level=2)
for item in [
    "Fund feasibility study: attractive central case, remains credible under downside testing, and no critical data gap.",
    "Monitor: potential value exists but the result is unstable or dependent on unresolved evidence.",
    "Do not fund: downside economics are weak or delivery constraints dominate likely benefits.",
    "Insufficient evidence: critical hazard, exposure, cost or intervention data is missing.",
]:
    add_bullet(doc, item)

add_callout(doc, "Governance principle", "The screening model recommends where to investigate—not where to build. Investment remains conditional on hydrological validation, ecological assessment, landowner agreement and a funded maintenance plan.", fill="FFF7E6", accent=GOLD)

doc.add_page_break()
doc.add_heading("8. Final presentation and application value", level=1)

doc.add_heading("Ten-minute interview presentation", level=2)
for item in [
    "The decision: why an insurer might invest in upstream wetland recovery.",
    "The pipeline: how inconsistent spatial datasets became controlled analytical inputs.",
    "The model: how hazard, exposure and vulnerability produce expected loss.",
    "The recommendation: which catchments merit feasibility work and why.",
    "The challenge: what uncertainty could reverse the decision.",
    "The reflection: what real claims and hydrological data would change next.",
]:
    add_bullet(doc, item)

doc.add_heading("Suggested CV statement", level=2)
add_callout(doc, "Portfolio summary", "Built a reproducible geospatial pipeline to identify where wetland restoration could reduce household flood exposure, combining hazard, property and habitat data to estimate avoided annual insured losses, benefit-cost ratios and model uncertainty.", fill=PALE_BLUE, accent=BLUE)

doc.add_heading("Final handover checklist", level=2)
for item in [
    "Repository runs from documented setup instructions.",
    "Sources, licences, retrieval dates and checksums are recorded.",
    "All critical data and calculation tests pass.",
    "Scenario assumptions are editable and version-controlled.",
    "Dashboard numbers reconcile to warehouse outputs.",
    "Recommendations distinguish observed data from assumptions.",
    "Model limitations and prohibited uses are prominent.",
    "Screenshots and committee paper contain no unexplained metrics.",
]:
    add_bullet(doc, item)

doc.add_heading("Starting references", level=2)
refs = [
    "Department for Environment, Food & Rural Affairs — Local Nature Recovery Strategies: https://www.gov.uk/government/publications/local-nature-recovery-strategies",
    "Natural England — Priority Habitats Inventory (England): https://www.data.gov.uk/dataset/4b6ddab7-6c0f-4407-946e-d6499f19fcde/priority-habitats-inventory-england",
    "Environment Agency — Flood and river data services: https://environment.data.gov.uk/",
]
for ref in refs:
    p = add_bullet(doc, ref)
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(MID)

doc.core_properties.title = "Quantifying the Insurance Value of Wetland Recovery"
doc.core_properties.subject = "Risk analytics portfolio project brief"
doc.core_properties.author = "Portfolio project specification"
doc.core_properties.keywords = "risk analytics, insurance, wetlands, flood risk, data engineering"

doc.save(OUT)
print(OUT)
