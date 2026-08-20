"""Week 6: a two-page executive briefing aimed at a flood-risk, insurance, or
nature-finance audience - a different tone and emphasis from the practitioner
outreach one-pager: less "here's my methodology," more "here's a screened
pipeline of candidate sites and what it would take to make them investable."

Rebuild with:
    python3 scripts/build_executive_briefing.py
"""
from pathlib import Path

import geopandas as gpd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
STAGING = ROOT / "data" / "processed" / "ribble_nfm_staging.gpkg"
LONGLIST = ROOT / "data" / "processed" / "candidate_longlist.gpkg"
MAPS_DIR = ROOT / "outputs" / "maps" / "week4"
OUT = ROOT / "outputs" / "Ribble_NFM_Executive_Briefing.docx"

FINAL_FIVE = ["C16", "C21", "C14", "C22", "C23"]

NAVY, BLUE, TEAL = "17324D", "2E74B5", "247D78"
INK, MUTED, WHITE = "243142", "667085", "FFFFFF"
LIGHT, PALE_BLUE, PALE_TEAL, GOLD = "F2F4F7", "E8EEF5", "E6F2F0", "8A6500"
RED, PALE_RED = "B71C1C", "FDECEA"


def font(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold, run.italic = bold, italic


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    pr.append(node)


def margins(cell, top=80, bottom=80, start=110, end=110):
    pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    pr.append(tc_mar)


def set_widths(table, widths_dxa):
    table.autofit = False
    tbl = table._tbl
    tblGrid = tbl.tblGrid
    for child in list(tblGrid):
        tblGrid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        tblGrid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def stat_row(doc, stats):
    """A row of headline-number tiles."""
    t = doc.add_table(rows=1, cols=len(stats))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (value, label) in enumerate(stats):
        c = t.cell(0, i)
        shade(c, PALE_BLUE)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        font(p.add_run(value), 18, NAVY, True)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p2.add_run(label), 8, MUTED)
    width = int(9360 / len(stats))
    set_widths(t, [width] * len(stats))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def callout(doc, label, text, fill=PALE_TEAL, accent=TEAL):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = t.cell(0, 0)
    shade(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run(label.upper()), 8.5, accent, True)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    font(p2.add_run(text), 9.6, NAVY, False)
    set_widths(t, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def bullet(doc, text, size=9.6):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(" " + text), size)
    return p


def section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(title), 12.5, BLUE, True)
    return p


def main():
    cands = gpd.read_file(LONGLIST)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(0.5)
    sec.left_margin = sec.right_margin = Inches(0.75)

    hp = sec.header.paragraphs[0]
    hp.text = "RIBBLE CATCHMENT NFM SCREENING  |  EXECUTIVE BRIEFING  |  SCREENING AREAS, NOT CONSTRUCTION BOUNDARIES"
    hp.paragraph_format.space_after = Pt(0)
    for r in hp.runs:
        font(r, 7.5, MUTED, True)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(fp.add_run("Pre-validation screening output  •  "), 8, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    fp._p.append(field)

    styles = doc.styles
    n = styles["Normal"]
    n.font.name, n.font.size = "Calibri", Pt(10)
    n.font.color.rgb = RGBColor.from_string(INK)
    n.paragraph_format.space_after = Pt(4)
    n.paragraph_format.line_spacing = 1.04
    s = styles["List Bullet"]
    s.font.name, s.font.size = "Calibri", Pt(9.6)
    s.paragraph_format.space_after = Pt(3)

    # ---------- PAGE 1 ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    font(p.add_run("GEOSPATIAL RISK SCREENING  —  RIBBLE CATCHMENT, LANCASHIRE"), 9, TEAL, True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    font(p.add_run("Five Candidate Natural Flood Management Sites: An Investment-Screening Pipeline"), 19, NAVY, True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run("A reproducible, data-led method for surfacing natural flood management opportunities that public "
                    "records don't yet show as delivered — built for flood-risk, insurance, and nature-finance audiences "
                    "evaluating where to focus due diligence, not where to build."), 10.5, MUTED, False, True)

    callout(doc, "Read this first", "Every site in this briefing is a SCREENING AREA identified from public data - "
            "not a proposed construction boundary, not an engineering design, and not investment-grade due diligence. "
            "It is a prioritisation tool for where to look next.", PALE_RED, RED)

    stat_row(doc, [
        ("1,404 km²", "Catchment screened"),
        ("35", "Candidates identified"),
        ("9", "Strong-band candidates"),
        ("282", "Recorded historic floods"),
        ("75", "Settlements in catchment"),
    ])

    section(doc, "The opportunity")
    p = doc.add_paragraph()
    font(p.add_run(
        "Natural flood management (NFM) - floodplain reconnection, wet woodland, leaky barriers, upland "
        "restoration - is increasingly relevant to flood-risk pricing, resilience planning, and nature-based "
        "finance, but identifying credible sites is slow, manual work. This project screens one catchment "
        "end-to-end against EA, Natural England, Historic England, and OS datasets, cross-checks modelled "
        "opportunity against recorded delivery, and screens out statutory protected-site constraints before a "
        "shortlist is even produced."
    ), 9.6)

    section(doc, "Why this matters, by audience")
    bullet(doc, "Flood risk: 22 of 35 candidates carry an explicit constraint flag (settlement or protected-site overlap) - a starting risk register, not a finished one.")
    bullet(doc, "Insurance / reinsurance: candidates are cross-referenced against 282 recorded historic flood events - a ground-truthed exposure signal, not a purely modelled one.")
    bullet(doc, "Nature finance: every candidate is scored on habitat-network overlap alongside flood potential, so dual flood/biodiversity value is visible from the outset.")

    # ---------- PAGE 2 ----------
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    section(doc, "The five candidates")
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, text in enumerate(["Site", "Nearest place", "Area", "Band", "Note for due diligence"]):
        shade(hdr[i], NAVY)
        pp = hdr[i].paragraphs[0]
        font(pp.add_run(text), 8.5, WHITE, True)
    rows = [
        ("C16", "Sabden", "300 ha", "Strong", "Cleanest candidate: no settlement or protected-site overlap, strong habitat-network overlap"),
        ("C21", "Barrowford", "187 ha", "Strong", "Clean, no visible constraints"),
        ("C14", "Withnell", "62 ha", "Strong", "Only strong candidate outside the Clitheroe/Pendle cluster - geographic diversification"),
        ("C22", "Gisburn", "600 ha", "Strong", "Clean, but rank is less stable across alternative scoring weights - treat with more caution"),
        ("C23", "Giggleswick", "571 ha", "Moderate", "Overlaps 2 SSSIs + a Scheduled Monument - likely needs statutory consent; included specifically to flag this"),
    ]
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            pp = cells[i].paragraphs[0]
            font(pp.add_run(val), 9)
    set_widths(t, [700, 1600, 900, 1100, 5060])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    section(doc, "How this was built (confidence basis)")
    bullet(doc, "8 national/regional datasets (EA, Natural England, Historic England, OS), all Open Government Licence, provenance-tracked to source and download date.")
    bullet(doc, "A six-factor scoring model with explicit, published weights, tested against 3 alternative weightings for rank stability - not a single black-box number.")
    bullet(doc, "A queryable SQL database (DuckDB) accompanies this briefing for independent verification of every figure quoted here.")

    section(doc, "What this is not")
    callout(doc, "Explicit limitations", "Desk-based public-data screening - not practitioner-validated, not "
            "hydrological or engineering assessment, and not a statement that any site is ready for capital "
            "deployment. Two of the five candidates carry explicit caveats (C22's rank stability, C23's "
            "protected-site overlap) to treat as open questions, not settled facts.",
            PALE_BLUE, BLUE)

    section(doc, "Recommended next steps")
    bullet(doc, "Local practitioner validation (in progress) with Ribble Rivers Trust - confirms or rejects each candidate against on-the-ground knowledge.")
    bullet(doc, "Statutory constraint check (SSSI/SAC consent pathway) specifically for C23 before any further spend.")
    bullet(doc, "Feasibility-grade hydrological and ecological assessment for any candidate carried past practitioner review - this screening explicitly does not replace that step.")

    doc.core_properties.title = "Ribble NFM Screening - Executive Briefing"
    doc.core_properties.subject = "Natural flood management investment-screening briefing"
    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
