"""Week 4: assemble the site-card document (catchment overview + opportunity-gap
map + five site cards, each with a map panel and an evidence panel), matching
the visual style of the original project overview.

Requires outputs/maps/week4/*.png (run scripts/build_week4_maps.py first).

Rebuild with:
    python3 scripts/build_site_cards.py
"""
from pathlib import Path

import geopandas as gpd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
STAGING = ROOT / "data" / "processed" / "ribble_nfm_staging.gpkg"
LONGLIST = ROOT / "data" / "processed" / "candidate_longlist.gpkg"
MAPS_DIR = ROOT / "outputs" / "maps" / "week4"
OUT = ROOT / "outputs" / "Ribble_NFM_Candidate_Site_Cards.docx"

FINAL_FIVE = ["C16", "C21", "C14", "C22", "C23"]

NAVY, BLUE, TEAL = "17324D", "2E74B5", "247D78"
INK, MUTED, WHITE = "243142", "667085", "FFFFFF"
LIGHT, PALE_BLUE, PALE_TEAL, GOLD = "F2F4F7", "E8EEF5", "E6F2F0", "8A6500"
AMBER = "B7791F"


def font(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold, run.italic = bold, italic


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    node = pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        pr.append(node)
    node.set(qn("w:fill"), fill)


def margins(cell, top=90, bottom=90, start=120, end=120):
    pr = cell._tc.get_or_add_tcPr()
    tc_mar = pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        el = tc_mar.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def geometry(table, widths):
    table.autofit = False
    pr = table._tbl.tblPr
    tw = pr.find(qn("w:tblW"))
    if tw is None:
        tw = OxmlElement("w:tblW")
        pr.append(tw)
    tw.set(qn("w:w"), str(sum(widths)))
    tw.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            pr = cell._tc.get_or_add_tcPr()
            tcw = pr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                pr.append(tcw)
            tcw.set(qn("w:w"), str(widths[i]))
            tcw.set(qn("w:type"), "dxa")
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def no_split_row(row):
    """Prevent this table row's content from splitting across a page
    boundary - if it doesn't fit, the whole row moves to the next page
    rather than being cut mid-row."""
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)


def evidence_table(doc, rows, widths, header_fill=NAVY):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in rows:
        tr = t.add_row()
        no_split_row(tr)
        r = tr.cells
        shade(r[0], PALE_BLUE)
        p = r[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(label.upper()), 8.6, NAVY, True)
        p2 = r[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        font(p2.add_run(value), 9.6, INK)
    geometry(t, widths)
    return t


def band_color(band):
    return {"Strong": TEAL, "Moderate": GOLD, "Watch": AMBER, "Weak": MUTED}.get(band, MUTED)


def callout(doc, label, text, fill=PALE_TEAL, accent=TEAL):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    geometry(t, [9360])
    c = t.cell(0, 0)
    shade(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(label.upper()), 9, accent, True)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    font(p2.add_run(text), 10.3, NAVY, False)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def infer_interventions(row):
    ideas = []
    if row["overlap_rofrs_pct"] > 1:
        ideas.append("floodplain reconnection / rewetting")
    if row["overlap_habitat_networks_pct"] > 50:
        ideas.append("wet woodland or habitat expansion")
    if row["dist_recorded_restoration_km"] > 1:
        ideas.append("runoff storage / leaky barriers")
    if not ideas:
        ideas.append("land-management measures (to be confirmed on site)")
    return "; ".join(ideas).capitalize()


def shorten_names(flag_text, max_names=2):
    """Keep evidence-table rows compact: a 3+ name protected-site list would
    wrap to several lines and be the main cause of a card overflowing one
    page. Full names stay in the DuckDB/CSV output; only the docx display
    is truncated."""
    names = [n.strip() for n in flag_text.split(";") if n.strip()]
    if len(names) <= max_names:
        return "; ".join(names)
    return "; ".join(names[:max_names]) + f" (+{len(names) - max_names} more - see docs/week4_manual_review.md)"


def build_evidence(row, nearest_name, nearest_dist_km):
    exposure = f"Nearest settlement: {nearest_name}, {nearest_dist_km:.2f}km away. {int(row['n_built_up_within_5km'])} settlements within 5km."
    if row["overlap_recorded_flood_pct"] > 0:
        exposure += f" {row['overlap_recorded_flood_pct']:.0f}% overlaps a recorded historic flood outline."
    habitat = f"{row['overlap_habitat_networks_pct']:.0f}% overlaps a mapped habitat-network opportunity zone; nearest existing priority habitat {row['dist_phi_km']:.2f}km away."
    if row["dist_recorded_restoration_km"] < 0.5:
        activity = f"A recorded restoration project sits within {row['dist_recorded_restoration_km']*1000:.0f}m - may already have known activity; verify locally."
    else:
        activity = f"Nearest recorded restoration project is {row['dist_recorded_restoration_km']:.2f}km away - no known nearby activity in the EA register."
    constraint_bits = []
    if row["overlap_built_up_pct"] > 5:
        constraint_bits.append(f"{row['overlap_built_up_pct']:.0f}% overlaps mapped settlement extent")
    if row["road_density_km_per_km2"] > 3:
        constraint_bits.append(f"road density {row['road_density_km_per_km2']:.1f}km/km2")
    protected_flag = row.get("protected_site_flag", "None")
    if protected_flag and str(protected_flag) != "None" and str(protected_flag).strip():
        constraint_bits.append(f"PROTECTED SITE: {shorten_names(str(protected_flag))} - likely needs consent")
    if constraint_bits:
        joined = "; ".join(constraint_bits)
        constraints = joined[0].upper() + joined[1:] + "."
    else:
        constraints = "No settlement or protected-site overlap; low road density."
    confidence = row["data_confidence"]
    next_check = "Practitioner walkover to confirm land use, ownership and access." if row["data_confidence"] == "High" else "Confirm flood-risk context locally - falls in the RoFRS acquisition gap."
    return dict(exposure=exposure, habitat=habitat, activity=activity, constraints=constraints, confidence=confidence, next_check=next_check)


def add_map_page(doc, title, image_path, caption, page_break_before=False):
    h = doc.add_heading(title, level=1)
    if page_break_before:
        h.paragraph_format.page_break_before = True
    doc.add_picture(str(image_path), width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    font(p.add_run(caption), 9, MUTED, False, True)


def add_site_card(doc, cid, row, nearest_name, nearest_dist_km, card_number, page_break_before=True):
    h = doc.add_heading(f"Site {card_number} — Candidate {cid}", level=1)
    h.paragraph_format.page_break_before = page_break_before
    color = band_color(row["band"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    font(p.add_run(f"{row['band']} candidate"), 12, color, True)
    font(p.add_run(f"   |   {row['area_ha']:.0f} ha   |   Rank {int(row['rank'])} of 35   |   Flood-risk context: {row['data_confidence']}"), 10.5, MUTED)

    doc.add_picture(str(MAPS_DIR / f"site_{cid}.png"), width=Inches(5.4))
    pic_para = doc.paragraphs[-1]
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.keep_with_next = True
    pic_para.paragraph_format.space_after = Pt(4)

    ev = build_evidence(row, nearest_name, nearest_dist_km)
    eh = doc.add_heading("Evidence", level=2)
    eh.paragraph_format.keep_with_next = True
    evidence_table(
        doc,
        [
            ("Location", f"Near {nearest_name}, Ribble catchment ({nearest_dist_km:.2f}km away)"),
            ("Interventions to investigate", infer_interventions(row)),
            ("Nearby community exposure", ev["exposure"]),
            ("Habitat value", ev["habitat"]),
            ("Recorded activity", ev["activity"]),
            ("Constraints", ev["constraints"]),
            ("Confidence (flood-risk context only)", ev["confidence"]),
            ("Recommended next check", ev["next_check"]),
        ],
        [2100, 7260],
    )


def main():
    cands = gpd.read_file(LONGLIST)
    bua = gpd.read_file(STAGING, layer="os_open_built_up_areas")

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(0.7)
    sec.left_margin = sec.right_margin = Inches(0.9)

    hp = sec.header.paragraphs[0]
    hp.text = "SCREENING AREAS — NOT PROPOSED CONSTRUCTION BOUNDARIES"
    hp.paragraph_format.space_after = Pt(0)
    for r in hp.runs:
        font(r, 8.5, "B71C1C", True)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(fp.add_run("Week 4 draft — desk-based, not practitioner-validated  •  "), 8.5, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    fp._p.append(field)

    styles = doc.styles
    n = styles["Normal"]
    n.font.name, n.font.size = "Calibri", Pt(11)
    n.font.color.rgb = RGBColor.from_string(INK)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
    ]:
        s = styles[name]
        s.font.name, s.font.size, s.font.bold = "Calibri", Pt(size), True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(before), Pt(after)
        s.paragraph_format.keep_with_next = True

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    font(p.add_run("WEEK 4 DRAFT — DESK-BASED REVIEW, NOT YET PRACTITIONER-VALIDATED"), 10, TEAL, True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    font(p.add_run("Ribble Catchment: Five Candidate NFM Sites"), 24, NAVY, True)

    callout(
        doc,
        "Read this before anything else",
        "Every polygon in this document is a SCREENING AREA showing where public data suggests further investigation may be worthwhile. None of them are proposed construction boundaries, engineering designs, or confirmed project sites - see docs/screening_methodology.md for what each score does and doesn't claim.",
        "FDECEA", "B71C1C",
    )

    callout(
        doc,
        "How these five were chosen",
        "From 35 modelled candidates, these five combine a strong or moderate composite score, rank stability across four alternative weightings, and catchment-wide geographic spread — after screening out oversized ‘Investigation Zones’ (≥ 1,000ha) and checking every candidate against SSSI, SAC and Scheduled Monument boundaries. One (C23) is kept specifically because it flags a real statutory constraint, not despite it. Full reasoning, including which higher-ranked candidates were set aside and why, is in docs/week4_manual_review.md.",
        PALE_BLUE, BLUE,
    )

    add_map_page(doc, "1. Catchment overview", MAPS_DIR / "01_catchment_overview.png",
                 "All 35 screened candidates, banded by composite score, with main rivers and settlements for context. The final five are outlined in navy.",
                 page_break_before=True)
    add_map_page(doc, "2. Opportunity-gap map", MAPS_DIR / "02_opportunity_gap.png",
                 "Modelled NFM potential (point size) against distance to the nearest recorded restoration project (colour). Green points are furthest from any recorded activity — the strongest 'gap' candidates.",
                 page_break_before=True)

    for i, cid in enumerate(FINAL_FIVE, start=1):
        row = cands[cands["candidate_id"] == cid].iloc[0]
        dists = bua.geometry.distance(row.geometry)
        nearest_idx = dists.idxmin()
        nearest_name = bua.loc[nearest_idx, "name1_text"]
        nearest_dist_km = dists.min() / 1000
        if i == 1:
            doc.add_heading("3. Site cards", level=1)
        # Every card forces its own fresh page - site 1 previously inherited
        # whatever room was left on the opportunity-gap map's page, which was
        # rarely a full page and caused it to overflow while cards 2-5 (each
        # starting fresh) did not.
        add_site_card(doc, cid, row, nearest_name, nearest_dist_km, i, page_break_before=True)

    doc.core_properties.title = "Ribble Catchment Candidate Site Cards"
    doc.core_properties.subject = "Natural flood management screening - Week 4 draft"
    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
