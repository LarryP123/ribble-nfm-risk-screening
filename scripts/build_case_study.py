"""Week 6 (draft): portfolio case study — architecture, controls, maps,
findings, and reflection, ready for an application or interview.

Status: pre-validation draft. The findings/reflection sections will be
revised once Week 5 practitioner feedback comes back — see docs/insight_briefing.md.

Rebuild with:
    python3 scripts/build_case_study.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
MAPS_DIR = ROOT / "outputs" / "maps" / "week4"
OUT = ROOT / "outputs" / "Ribble_NFM_Portfolio_Case_Study.docx"

NAVY, BLUE, TEAL = "17324D", "2E74B5", "247D78"
INK, MUTED, WHITE = "243142", "667085", "FFFFFF"
LIGHT, PALE_BLUE, PALE_TEAL, GOLD = "F2F4F7", "E8EEF5", "E6F2F0", "8A6500"


def font(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold, run.italic = bold, italic


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    pr.append(node)


def margins(cell, top=90, bottom=90, start=120, end=120):
    pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    pr.append(tc_mar)


def geometry(table, widths):
    table.autofit = False
    pr = table._tbl.tblPr
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(sum(widths)))
    tw.set(qn("w:type"), "dxa")
    pr.append(tw)
    grid = table._tbl.tblGrid
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cpr = cell._tc.get_or_add_tcPr()
            tcw = OxmlElement("w:tcW")
            tcw.set(qn("w:w"), str(widths[i]))
            tcw.set(qn("w:type"), "dxa")
            cpr.append(tcw)
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(" " + text)
    return p


def table(doc, headers, rows, widths, size=8.9, header_fill=NAVY):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, text in enumerate(headers):
        shade(t.rows[0].cells[i], header_fill)
        p = t.rows[0].cells[i].paragraphs[0]
        font(p.add_run(text), size, WHITE, True)
    for ri, values in enumerate(rows):
        cells = t.add_row().cells
        if ri % 2:
            for c in cells:
                shade(c, LIGHT)
        for i, text in enumerate(values):
            p = cells[i].paragraphs[0]
            font(p.add_run(str(text)), size)
    geometry(t, widths)
    doc.add_paragraph()
    return t


def callout(doc, label, text, fill=PALE_TEAL, accent=TEAL):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    geometry(t, [9360])
    c = t.cell(0, 0)
    shade(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(label.upper()), 9, accent, True)
    p2 = c.add_paragraph()
    font(p2.add_run(text), 10.3, NAVY, False)
    doc.add_paragraph()


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.75)
sec.left_margin = sec.right_margin = Inches(0.9)

hp = sec.header.paragraphs[0]
hp.text = "PORTFOLIO CASE STUDY  |  RIBBLE NFM SCREENING"
for r in hp.runs:
    font(r, 8.5, MUTED, True)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(fp.add_run("Draft — pending Week 5 validation  •  "), 8.5, MUTED)
field = OxmlElement("w:fldSimple")
field.set(qn("w:instr"), "PAGE")
fp._p.append(field)

styles = doc.styles
n = styles["Normal"]
n.font.name, n.font.size = "Calibri", Pt(11)
n.font.color.rgb = RGBColor.from_string(INK)
n.paragraph_format.space_after = Pt(5)
n.paragraph_format.line_spacing = 1.08
for name, size, color, before, after in [("Heading 1", 15, BLUE, 12, 6), ("Heading 2", 12.5, BLUE, 9, 5)]:
    s = styles[name]
    s.font.name, s.font.size, s.font.bold = "Calibri", Pt(size), True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(before), Pt(after)
    s.paragraph_format.keep_with_next = True
for name in ("List Bullet",):
    s = styles[name]
    s.font.name, s.font.size = "Calibri", Pt(11)
    s.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
font(p.add_run("GEOSPATIAL RISK PORTFOLIO PROJECT — CASE STUDY"), 10, TEAL, True)
p = doc.add_paragraph()
font(p.add_run("Finding Overlooked Natural Flood-Management Sites in the Ribble Catchment"), 24, NAVY, True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
font(p.add_run("A reproducible spatial pipeline for candidate NFM site screening, built and documented end-to-end"), 12.5, MUTED)

callout(doc, "Status", "Draft — Weeks 1-4 complete (methodology corrected once, see Reflection), Week 5 outreach materials prepared for Ribble Rivers Trust but not yet sent. Findings and reflection below will be revised again once practitioner feedback lands (see docs/insight_briefing.md).", PALE_BLUE, BLUE)

doc.add_heading("1. Purpose", level=1)
body(doc, "Combine national flood, habitat, and restoration datasets to identify candidate natural flood management areas in one catchment, then test the shortlist against local practitioner knowledge to separate genuine gaps from artefacts of incomplete public data.")

doc.add_heading("2. Architecture", level=1)
body(doc, "A five-stage pipeline, each stage a documented, rerunnable script reading only from the stage before it:")
table(doc, ["STAGE", "SCRIPT", "OUTPUT"], [
    ("1. Acquisition", "Manual + scripts/fetch_ogc_features.py", "7 raw datasets in data/raw/ (national + catchment-scale)"),
    ("2. Staging", "scripts/build_staging_layers.py", "One spatial database: 15 layers, common CRS, exact-catchment clip"),
    ("3. Scoring", "scripts/score_candidates.py", "35 candidate sites, 6-factor score, sensitivity-checked ranking"),
    ("4. Map production", "scripts/build_week4_maps.py, build_site_cards.py", "Catchment/gap/site maps + assembled site-card document"),
    ("5. Outreach", "scripts/build_outreach_one_pager.py", "One-page summary + map for practitioner review"),
], [1600, 3600, 4160])

doc.add_heading("3. Controls", level=1)
for text in [
    "Provenance manifest (docs/provenance_manifest.csv): every dataset's source, licence, download date, and known gaps, including ones that were hard-won or incomplete.",
    "Data dictionary (docs/data_dictionary.md): every staging layer's schema, feature count, and validation result — 0 invalid/empty geometries, consistent CRS across all 15 layers.",
    "Explicit, documented scoring weights, tested across 4 alternative weightings for rank stability, rather than presented as a single precise answer.",
    "A written manual-review record (docs/week4_manual_review.md) showing which higher-ranked candidates were set aside and why, not just which were kept.",
]:
    bullet(doc, text)

doc.add_heading("4. Key maps", level=1)
doc.add_picture(str(MAPS_DIR / "01_catchment_overview.png"), width=Inches(3.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
font(p.add_run("All 35 screened candidates, banded by score, with the final five outlined."), 9, MUTED, False, True)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()
doc.add_heading("5. Findings", level=1)
body(doc, "35 candidates were identified from 41 raw modelled opportunity polygons; any candidate 1,000ha or larger is labelled an Investigation Zone rather than a site. Five sites were carried forward as site cards (C16 Sabden, C21 Barrowford, C14 Withnell, C22 Gisburn, C23 Giggleswick), chosen for composite score, rank stability across weightings, and deliberate catchment-wide spread rather than raw rank alone.")
callout(doc, "A genuine pattern, not an artefact", "The top-ranked, rank-stable candidates cluster heavily in the Pendle/Ribble Valley area near Clitheroe. The modelled NFM signal in this catchment is unevenly distributed - worth stating plainly rather than presenting five geographically arbitrary points.", PALE_BLUE, BLUE)
body(doc, "One site (C23) was kept despite a real flagged constraint - it overlaps two SSSIs and a Scheduled Monument - specifically because surfacing that flag for practitioner review is more useful than quietly filtering it out. A different candidate, C17, was dropped entirely after a correction found it was both oversized (2,076ha - an Investigation Zone, not a site) and overlapped five separate SSSIs; see Reflection.")

doc.add_heading("6. Reflection", level=1)
body(doc, "Four moments in this project changed the outcome, worth being specific about rather than summarised away:")
for label, text in [
    ("The RoFRS acquisition had no clean path.",
     "Unlike every other dataset, RoFRS exposed no bbox/vector API - only WMS tiles and an area-of-interest export tool that rejected large polygons. Reaching 98.8% coverage took 12 manual exports, repeated checks against the real catchment boundary, and fixing a pagination bug - the API's cursor was startIndex, not the offset I assumed, which silently returned the same page twice until a spot-check caught it."),
    ("The Week 1 data list was checked against the brief and found short.",
     "Re-reading the original brief's data table against what had actually been downloaded showed two of seven planned sources - land cover/roads/settlements, and rivers/terrain - had been dropped without being flagged, needed for two of six scoring factors. Catching this meant treating it as a real gap, not working around it with weaker proxies."),
    ("Manual review isn't mechanical top-N selection.",
     "Taking the top five by score alone would have produced four sites within 8km of each other and one outlier - technically correct, much less useful for a practitioner conversation. The final selection traded some rank for spread, and that trade-off is written down, not hidden inside a black-box score."),
    ("A later review found the scoring itself was incomplete, and it changed the answer.",
     "A closer read of the methodology - not practitioner feedback - found RoFRS overlap was computed but never used or explained, recorded flood outlines weren't used at all, and no protected-site data existed, so a 2,076ha candidate with five SSSI overlaps (C17) had sat at rank 2 undetected. Rescoring changed the final five - C17 out, C22 in. A high, stable rank is not the same as a checked one."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    font(p.add_run(label + " "), 10.3, NAVY, True)
    font(p.add_run(text), 10.3, INK)

body(doc, "What I'd still do differently: build proper flow-direction analysis for community exposure instead of settlement proximity, and get real sub-parcel data to responsibly split the Investigation Zones instead of just flagging their size.")

doc.core_properties.title = "Ribble NFM Screening - Portfolio Case Study"
doc.core_properties.subject = "Geospatial risk portfolio project - case study"
doc.save(OUT)
print(f"Written: {OUT}")
