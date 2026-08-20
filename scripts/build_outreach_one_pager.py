"""Week 5: one-page outreach summary for a Ribble Rivers Trust contact,
per the brief's "single-page summary and one strong map" outreach approach.

This is deliberately terse - see docs/query_database.md and the case study
for the full picture. The point of this document is to fit on one page.

Rebuild with:
    python3 scripts/build_outreach_one_pager.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
MAPS_DIR = ROOT / "outputs" / "maps" / "week4"
OUT = ROOT / "outputs" / "Ribble_Outreach_One_Pager.docx"

NAVY, BLUE, TEAL = "17324D", "2E74B5", "247D78"
INK, MUTED, WHITE = "243142", "667085", "FFFFFF"
PALE_BLUE = "E8EEF5"


def font(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold, run.italic = bold, italic


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    pr.append(node)


def tight(p, space_after=3, space_before=0):
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0


def set_col_widths(table, widths_dxa):
    table.autofit = False
    tbl = table._tbl
    tblGrid = tbl.tblGrid
    for child in list(tblGrid):
        tblGrid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        tblGrid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.55)
sec.left_margin = sec.right_margin = Inches(0.7)

styles = doc.styles
n = styles["Normal"]
n.font.name, n.font.size = "Calibri", Pt(9.5)
n.font.color.rgb = RGBColor.from_string(INK)
n.paragraph_format.space_after = Pt(3)
n.paragraph_format.line_spacing = 1.0

# --- Header block ---
p = doc.add_paragraph(); tight(p, 2)
font(p.add_run("RIBBLE CATCHMENT — NATURAL FLOOD MANAGEMENT SCREENING"), 8.5, TEAL, True)
p = doc.add_paragraph(); tight(p, 2)
font(p.add_run("Five candidate sites from a data-led screening exercise"), 16, NAVY, True)
p = doc.add_paragraph(); tight(p, 6)
font(p.add_run("A screening tool, not a proposal — every polygon is a SCREENING AREA, not a proposed construction boundary."), 9, "B71C1C", True)

# --- Intro (short) + map side by side via a borderless table ---
layout = doc.add_table(rows=1, cols=2)
cell_l, cell_r = layout.rows[0].cells

p = cell_l.paragraphs[0]; tight(p, 4)
font(p.add_run(
    "I've combined EA NFM Heat Maps, flood-risk and recorded-flooding data, Natural England habitat/SSSI/SAC "
    "layers, Historic England Scheduled Monuments, and OS settlement/river data to screen the catchment for "
    "strong modelled NFM potential with little recorded delivery — genuine gaps, not sites already known, "
    "worked, or statutorily constrained."
), 9.5)

p = cell_l.add_paragraph(); tight(p, 2, 6)
font(p.add_run("What I'd value from you"), 10.5, BLUE, True)
for text in [
    "Do these five reflect real conditions, or is there unrecorded activity/constraints public data wouldn't show?",
    "Any location you'd expect flagged that isn't here?",
    "Does C23's SSSI/Monument overlap rule it out, or is a smaller area within it still feasible?",
]:
    p = cell_l.add_paragraph(style="List Bullet"); tight(p, 2)
    font(p.add_run(text), 9)

cell_r.paragraphs[0].add_run().add_picture(str(MAPS_DIR / "03_outreach_locator.png"), width=Inches(3.3))
cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

set_col_widths(layout, [3900, 3300])
for cell in (cell_l, cell_r):
    tcPr = cell._tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        tcBorders.append(el)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# --- Table of five ---
p = doc.add_paragraph(); tight(p, 3, 4)
font(p.add_run("Five candidates, chosen for strength and catchment-wide spread"), 11, BLUE, True)

t = doc.add_table(rows=1, cols=4)
t.style = "Table Grid"
hdr = t.rows[0].cells
for i, text in enumerate(["Site", "Nearest place", "Area", "Note"]):
    shade(hdr[i], NAVY)
    pp = hdr[i].paragraphs[0]; tight(pp, 1)
    font(pp.add_run(text), 8.5, WHITE, True)

rows = [
    ("C16", "Sabden", "300 ha", "Top-ranked; no settlement or protected-site overlap; strong habitat-network overlap"),
    ("C21", "Barrowford", "187 ha", "Clean, no visible constraints"),
    ("C14", "Withnell", "62 ha", "Only strong candidate outside the Clitheroe/Pendle cluster"),
    ("C22", "Gisburn", "600 ha", "Clean, but rank is less stable across alternative weightings than the others"),
    ("C23", "Giggleswick", "571 ha", "Only credible candidate near the upper catchment - overlaps two SSSIs and a scheduled monument, likely needs consent. Kept deliberately, not hidden"),
]
for r in rows:
    cells = t.add_row().cells
    for i, val in enumerate(r):
        pp = cells[i].paragraphs[0]; tight(pp, 1)
        font(pp.add_run(val), 8.5)

set_col_widths(t, [650, 1500, 750, 6924])

p = doc.add_paragraph(); tight(p, 0, 6)
font(p.add_run(
    "This narrows the search and makes uncertainty visible; it doesn't claim these are ready for investment or "
    "that public data is complete. Happy to talk through the maps, or take feedback by email."
), 9, INK, False, True)

doc.core_properties.title = "Ribble Catchment NFM Screening - Outreach Summary"
doc.save(OUT)
print(f"Written: {OUT}")
